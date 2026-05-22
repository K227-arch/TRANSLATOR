"""
generate_data_report.py
=======================
Generates a Word document report covering:
1. What was cleaned from runyoro_dictionary_with_domains.xlsx
2. Domain distribution across ALL combined datasets
"""

import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = Path(__file__).parent
DATA_DIR = BASE / "data"
OUT_DOC  = BASE / "data" / "word" / "dataset_cleaning_and_domain_report.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F4E79"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, header_bg="1F4E79", alt_bg="D6E4F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if c_idx > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


# ── 1. Cleaning stats (re-derive from source) ─────────────────────────────────

def get_cleaning_stats():
    import re

    xl = pd.ExcelFile(DATA_DIR / "raw" / "runyoro_dictionary_with_domains.xlsx")

    all_rows = []
    for sheet in xl.sheet_names[1:]:
        df = pd.read_excel(xl, sheet_name=sheet, header=None)
        header_row = None
        for i, row in df.iterrows():
            vals = [str(v).strip().lower() for v in row.values]
            if 'word' in vals and any('definition' in v or 'meaning' in v for v in vals):
                header_row = i
                break
        if header_row is None:
            continue
        df.columns = df.iloc[header_row]
        df = df.iloc[header_row + 1:].reset_index(drop=True)
        df.columns = [str(c).strip() for c in df.columns]
        col_map = {}
        for c in df.columns:
            cl = c.lower()
            if 'word' in cl and 'count' not in cl: col_map.setdefault('word', c)
            elif 'definition' in cl or 'meaning' in cl: col_map.setdefault('definition', c)
            elif 'domain' in cl: col_map.setdefault('domain', c)
            elif 'part' in cl or 'speech' in cl: col_map.setdefault('pos', c)
        if 'word' not in col_map or 'definition' not in col_map:
            continue
        cols = [col_map['word'], col_map['definition']]
        if 'domain' in col_map: cols.append(col_map['domain'])
        sub = df[cols].copy()
        sub.columns = ['word', 'definition'] + (['domain'] if 'domain' in col_map else [])
        sub['sheet'] = sheet
        all_rows.append(sub)

    raw = pd.concat(all_rows, ignore_index=True)
    raw['word']       = raw['word'].astype(str).str.strip()
    raw['definition'] = raw['definition'].astype(str).str.strip()
    n_raw = len(raw)

    steps = []

    # Step 1: header rows
    mask = (raw['word'].str.lower().isin(['word','speech','nan','']) |
            raw['definition'].str.lower().isin(['english definition','meaning/definition','nan','']))
    steps.append(("Leaked header rows", int(mask.sum()), "Rows where word='Word' or definition='English Definition' leaked from Excel headers"))
    raw = raw[~mask].copy()

    # Step 2: morpheme fragments
    mask = raw['word'].str.startswith('-') | raw['word'].str.endswith('-')
    steps.append(("Morpheme fragments (starts/ends with -)", int(mask.sum()), "Grammatical prefixes/suffixes like -a-, oku-, -ire that are not standalone words"))
    raw = raw[~mask].copy()

    # Step 3: single/two-char entries
    mask = raw['word'].str.len() <= 2
    steps.append(("Single/two-character entries", int(mask.sum()), "Entries like 'a', 'o-', 'bi' that are particles or prefixes, not translatable words"))
    raw = raw[~mask].copy()

    # Step 4: grammar notation definitions
    RE_GRAMMAR = re.compile(r'Possess\.\s*particle|pronom\.\s*prefix|subj\.\s*pronom|tense\s+prefix|formative|concord\s+in\s+use', re.I)
    mask = raw['definition'].str.contains(RE_GRAMMAR, na=False)
    steps.append(("Grammar-notation definitions", int(mask.sum()), "Definitions describing grammatical function (e.g. 'Possess. particle, ro which obj. rel.cone...')"))
    raw = raw[~mask].copy()

    # Step 5: abbreviation-only
    mask = raw['definition'].str.match(r'^[a-z]{1,5}\.\s*(cl\.|v\.|adj\.|n\.)', na=False)
    steps.append(("Abbreviation-only definitions", int(mask.sum()), "Definitions that are just abbreviations like 'n. cl. 5' with no actual meaning"))
    raw = raw[~mask].copy()

    # Step 6: too short
    mask = raw['definition'].str.len() < 4
    steps.append(("Definitions shorter than 4 characters", int(mask.sum()), "Entries with no meaningful English content"))
    raw = raw[~mask].copy()

    # Step 7: OCR fixes (count affected)
    OCR_FIXES = [('RunyoroRutooro','Runyoro-Rutooro'),('wh ich','which'),('poessor','possessor'),('cone\\.','concord.'),('perfonn','perform')]
    ocr_count = 0
    for pat, rep in OCR_FIXES:
        before = raw['definition'].copy()
        raw['definition'] = raw['definition'].str.replace(pat, rep, regex=True)
        ocr_count += int((raw['definition'] != before).sum())
    steps.append(("OCR/typo artifacts fixed", ocr_count, "Corrected: RunyoroRutooro→Runyoro-Rutooro, 'wh ich'→'which', 'poessor'→'possessor', etc."))

    # Step 8: leading grammar labels stripped
    RE_LABEL = re.compile(r'^(?:n\.|v\.|adj\.|adv\.|pron\.|prep\.|conj\.|int\.|nt\.|j\.|part\.|interj\.|num\.|art\.)\s+', re.I)
    before = raw['definition'].copy()
    raw['definition'] = raw['definition'].str.replace(RE_LABEL, '', regex=True).str.strip()
    steps.append(("Leading grammar labels stripped", int((raw['definition'] != before).sum()), "Removed prefixes like 'n. ', 'v. ', 'adj. ' from the start of definitions"))

    # Step 9: trailing punctuation
    before = raw['definition'].copy()
    raw['definition'] = raw['definition'].str.rstrip(' .,;:').str.replace(r'\s+', ' ', regex=True).str.strip()
    steps.append(("Trailing punctuation noise fixed", int((raw['definition'] != before).sum()), "Removed trailing dots, commas, semicolons and extra whitespace"))

    # Step 10: domain assignment
    no_domain = raw.get('domain', pd.Series([''] * len(raw))).astype(str).str.lower().isin(['nan','','none'])
    steps.append(("Entries assigned 'General' domain", int(no_domain.sum()), "Entries with no domain tag were assigned the 'General' domain"))

    # Step 11: duplicates
    before_len = len(raw)
    raw = raw.drop_duplicates(subset=['word','definition'], keep='first')
    steps.append(("Duplicate word+definition pairs removed", before_len - len(raw), "Identical word+definition combinations kept only once"))

    n_after_removal = sum(s[1] for s in steps[:6])  # actual removals
    n_final = n_raw - n_after_removal

    # Split: direct vs definitions
    raw['wc'] = raw['definition'].str.split().str.len()
    n_direct = int((raw['wc'] <= 4).sum())
    n_defs   = int((raw['wc'] >  4).sum())

    return steps, n_raw, n_final, n_direct, n_defs


# ── 2. Domain distribution across ALL datasets ────────────────────────────────

def get_domain_distribution():
    domain_counts = {}

    def add_domain(label, count):
        domain_counts[label] = domain_counts.get(label, 0) + count

    # --- runyoro_domain_dictionary_clean.csv (direct translations) ---
    d = pd.read_csv(DATA_DIR / "cleaned" / "runyoro_domain_dictionary_clean.csv")
    for dom, cnt in d['domain'].value_counts().items():
        add_domain(str(dom).strip(), int(cnt))

    # --- dictionary_lookup.csv (definitions) ---
    l = pd.read_csv(DATA_DIR / "cleaned" / "dictionary_lookup.csv")
    for dom, cnt in l['domain'].value_counts().items():
        add_domain(str(dom).strip(), int(cnt))

    # --- train.csv domain-tagged pairs ---
    train = pd.read_csv(DATA_DIR / "training" / "train.csv")
    tagged = train[train['english'].str.match(r'^\[[A-Z_]+\]', na=False)]
    extracted = tagged['english'].str.extract(r'^\[([^\]]+)\]')[0]
    for dom, cnt in extracted.value_counts().items():
        # Normalise to title case
        dom_clean = str(dom).replace('_', ' ').title()
        add_domain(dom_clean, int(cnt))

    # --- seed vocabulary files ---
    seed_files = {
        'Health & Medicine':  DATA_DIR / "raw" / "medical_seed_vocabulary.csv",
        'Agriculture':        DATA_DIR / "raw" / "agriculture_seed_vocabulary.csv",
        'Education':          DATA_DIR / "raw" / "education_seed_vocabulary.csv",
        'Daily Life':         DATA_DIR / "raw" / "daily_life_seed_vocabulary.csv",
    }
    for dom, path in seed_files.items():
        if path.exists():
            df = pd.read_csv(path)
            add_domain(dom, len(df))

    # --- cleaned sources (no domain tag → General) ---
    untagged_sources = {
        'General': [
            DATA_DIR / "cleaned" / "english_nyoro_clean.csv",
            DATA_DIR / "cleaned" / "runyoro_english_sentences_clean.csv",
            DATA_DIR / "cleaned" / "word_entries_clean.csv",
            DATA_DIR / "cleaned" / "rutooro_dictionary_clean.csv",
            DATA_DIR / "cleaned" / "ocr_pairs_extracted.csv",
            DATA_DIR / "cleaned" / "ocr_glosses_extracted.csv",
        ],
        'Culture & Tradition': [DATA_DIR / "cleaned" / "empaako_pairs.csv"],
        'Proverbs':            [DATA_DIR / "cleaned" / "proverbs_pairs_clean.csv"],
        'Interjections':       [DATA_DIR / "cleaned" / "interjections_pairs_clean.csv"],
        'Numbers':             [DATA_DIR / "cleaned" / "numbers_pairs.csv"],
        'Idioms':              [DATA_DIR / "cleaned" / "idioms_pairs.csv"],
        'Grammar (Gr4)':       [DATA_DIR / "cleaned" / "gr4_pairs.csv"],
        'Grammar (Gr5)':       [DATA_DIR / "cleaned" / "gr5_pairs.csv"],
    }
    for dom, paths in untagged_sources.items():
        for path in paths:
            if path.exists():
                df = pd.read_csv(path)
                add_domain(dom, len(df))

    # Normalise duplicate domain names
    MERGE_MAP = {
        'Health And Medicine':           'Health & Medicine',
        'Nature And Environment':        'Nature & Environment',
        'Culture And Tradition':         'Culture & Tradition',
        'Education And Learning':        'Education & Learning',
        'Education':                     'Education & Learning',
        'Governance And Administration': 'Governance & Administration',
        'Art And Entertainment':         'Art & Entertainment',
        'Legal And Law Matters':         'Legal & Law',
        'Economics And Commerce':        'Economics & Commerce',
        'Politics And Current Affairs':  'Politics & Current Affairs',
        'Storytelling And Narratives':   'Storytelling & Narratives',
        'Astronomy And The Universe':    'Astronomy & Universe',
        'History And Historical Accounts': 'History',
    }
    merged = {}
    for dom, cnt in domain_counts.items():
        canonical = MERGE_MAP.get(dom, dom)
        merged[canonical] = merged.get(canonical, 0) + cnt

    total = sum(merged.values())
    rows = sorted(merged.items(), key=lambda x: -x[1])
    return rows, total


# ── Build Word document ───────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading('Runyoro-Rutooro Translator', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string("1F4E79")

    sub = doc.add_paragraph('Dataset Cleaning & Domain Distribution Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor.from_string("2E75B6")

    date_p = doc.add_paragraph(f'Generated: {datetime.now().strftime("%d %B %Y, %H:%M")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor.from_string("808080")

    doc.add_paragraph()

    # ── Section 1: Source document overview ───────────────────────────────────
    add_heading(doc, '1. Source Document Overview', level=1)
    doc.add_paragraph(
        'The new document added to the raw data folder is runyoro_dictionary_with_domains.xlsx — '
        'a comprehensive Runyoro-Rutooro dictionary organised alphabetically across 24 sheets '
        '(one per letter, plus a Summary sheet). Each entry contains a Runyoro/Rutooro word, '
        'its English definition, part of speech, and an optional domain tag.'
    )

    add_heading(doc, 'Document Structure', level=2)
    add_table(doc,
        ['Property', 'Value'],
        [
            ['File name',         'runyoro_dictionary_with_domains.xlsx'],
            ['Sheets',            '24 (1 Summary + 23 letter sheets: A–Z excl. H, V, X)'],
            ['Total raw entries', '21,512'],
            ['Columns per entry', 'Word, Part of Speech, English Definition, Domain'],
            ['Largest letter',    'O — 7,060 entries'],
            ['Domains present',   '23 named domains + untagged entries'],
        ]
    )
    doc.add_paragraph()

    # ── Section 2: Cleaning report ────────────────────────────────────────────
    add_heading(doc, '2. Cleaning Steps Applied', level=1)
    doc.add_paragraph(
        'The following cleaning pipeline was applied to the raw dictionary data. '
        'Steps 1–6 remove rows entirely; steps 7–11 fix or enrich existing data.'
    )

    steps, n_raw, n_final, n_direct, n_defs = get_cleaning_stats()

    rows = []
    for i, (name, count, reason) in enumerate(steps, 1):
        action = 'Removed' if i <= 6 else 'Fixed/Assigned'
        rows.append([f'{i}. {name}', f'{count:,}', action, reason])

    add_table(doc,
        ['Cleaning Step', 'Count', 'Action', 'Description'],
        rows
    )
    doc.add_paragraph()

    # Summary box
    add_heading(doc, 'Cleaning Summary', level=2)
    add_table(doc,
        ['Metric', 'Count'],
        [
            ['Raw entries loaded',                    f'{n_raw:,}'],
            ['Entries removed (steps 1–6)',           f'{sum(s[1] for s in steps[:6]):,}'],
            ['Entries after removal',                 f'{n_raw - sum(s[1] for s in steps[:6]):,}'],
            ['OCR/formatting fixes applied',          f'{sum(s[1] for s in steps[6:9]):,}'],
            ['Entries assigned General domain',       f'{steps[9][1]:,}'],
            ['Duplicates removed',                    f'{steps[10][1]:,}'],
            ['FINAL clean entries',                   f'{n_final:,}'],
            ['→ Direct translations (≤4 words)',      f'{n_direct:,}  → runyoro_domain_dictionary_clean.csv'],
            ['→ Definitions/explanations (>4 words)', f'{n_defs:,}  → dictionary_lookup.csv'],
        ]
    )
    doc.add_paragraph()

    # ── Section 3: Why the split ──────────────────────────────────────────────
    add_heading(doc, '3. Direct Translations vs Definitions', level=1)
    doc.add_paragraph(
        'The English text in this dictionary is not always a direct translation — '
        'in many cases it is a descriptive definition or explanation of the Runyoro word. '
        'Using definitions as MT training pairs would teach the model to output explanations '
        'instead of concise translations. Therefore the data was split:'
    )
    add_table(doc,
        ['Type', 'Criterion', 'Count', 'Destination', 'Purpose'],
        [
            ['Direct translation', '≤ 4 English words', f'{n_direct:,}',
             'runyoro_domain_dictionary_clean.csv\n+ train.csv / val.csv',
             'MT model training pairs'],
            ['Definition/explanation', '> 4 English words', f'{n_defs:,}',
             'dictionary_lookup.csv',
             'Dictionary fallback lookup in translate.py'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Example direct translation: okubura → to lose\n'
        'Example definition: ekituho → stench: a horrid and offensive smell coming out of a decaying substance'
    ).runs[0].font.italic = True

    # ── Section 4: Domain distribution ───────────────────────────────────────
    add_heading(doc, '4. Domain Distribution Across All Datasets', level=1)
    doc.add_paragraph(
        'The following table shows the domain distribution across all combined data sources: '
        'training pairs (train.csv), the domain dictionary, the definition lookup, '
        'seed vocabulary files, and all cleaned source CSVs.'
    )

    domain_rows, total = get_domain_distribution()

    table_rows = []
    for rank, (dom, cnt) in enumerate(domain_rows, 1):
        pct = 100 * cnt / total
        bar = '█' * int(pct / 2) if pct >= 1 else '▌' if pct >= 0.5 else '·'
        table_rows.append([str(rank), dom, f'{cnt:,}', f'{pct:.2f}%', bar])

    table_rows.append(['', 'TOTAL', f'{total:,}', '100.00%', ''])

    add_table(doc,
        ['#', 'Domain', 'Entries', 'Percentage', 'Visual'],
        table_rows
    )
    doc.add_paragraph()

    # ── Section 5: Training data impact ──────────────────────────────────────
    add_heading(doc, '5. Training Data Impact', level=1)
    add_table(doc,
        ['Dataset', 'Before', 'After', 'Added'],
        [
            ['train.csv', '67,939', '76,010', '+8,071'],
            ['val.csv',   '4,264',  '5,161',  '+897'],
            ['Total',     '72,203', '81,171', '+8,968'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The 8,968 direct translation pairs were added with domain prefix tags '
        '(e.g. [GENERAL], [COMMON_ANIMALS], [HEALTH_AND_MEDICINE]) to enable '
        'domain-aware translation. The 12,059 definition entries are available '
        'to the dictionary fallback system in translate.py.'
    )

    # ── Save ──────────────────────────────────────────────────────────────────
    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOC)
    print(f'Report saved: {OUT_DOC}')
    return domain_rows, total


if __name__ == '__main__':
    print('Generating report...')
    domain_rows, total = build_report()
    print(f'\nDomain distribution ({len(domain_rows)} domains, {total:,} total entries):')
    for dom, cnt in domain_rows:
        pct = 100 * cnt / total
        print(f'  {dom:<45} {cnt:>8,}  ({pct:.2f}%)')
