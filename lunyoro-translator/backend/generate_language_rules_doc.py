"""
generate_language_rules_doc.py
================================
Generates a Word document explaining how all language rules are applied
in the Runyoro-Rutooro translator post-processing pipeline.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path

OUT = Path(__file__).parent / "data" / "word" / "language_rules_explanation.docx"

NAVY   = "1F4E79"
BLUE   = "2E75B6"
TEAL   = "00695C"
ORANGE = "C55A11"
GRAY   = "595959"
WHITE  = "FFFFFF"
LIGHT  = "D6E4F0"
ALT    = "EBF5FB"


def shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_color)
    tcPr.append(s)


def heading(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def para(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def table(doc, headers, rows, hbg=NAVY, alt=LIGHT):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        shd(c, hbg)
        run = c.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = alt if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            shd(c, bg)
            c.paragraphs[0].runs[0].font.size = Pt(9)
    return t


def example_box(doc, runyoro, english, note=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"{runyoro}  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(TEAL)
    r1.font.size = Pt(10)
    r2 = p.add_run(f"→  {english}")
    r2.font.size = Pt(10)
    if note:
        r3 = p.add_run(f"  [{note}]")
        r3.italic = True
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor.from_string(GRAY)


def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    t = doc.add_heading("Runyoro-Rutooro Translator", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.color.rgb = RGBColor.from_string(NAVY)

    s = doc.add_paragraph("Language Rules — How They Are Applied")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(14)
    s.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    d = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.runs[0].font.size = Pt(10)
    d.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph()

    # ── Overview ──────────────────────────────────────────────────────────────
    heading(doc, "1. Overview", 1)
    para(doc,
        "The translator applies a layered post-processing pipeline to every "
        "machine-translated Runyoro-Rutooro output. Rules are applied in a fixed "
        "order because some rules depend on the output of earlier ones. "
        "The pipeline is defined in translate.py → _postprocess_lunyoro() and "
        "calls functions from three rule modules:", size=10)
    table(doc,
        ["Module", "Source", "Rules covered"],
        [
            ["language_rules.py",     "Grammar Ch.2,4,7,13,15,16 + Orthography Guide 1995",
             "R/L rule, nasal assimilation, ni-prefix, apostrophe elision, semi-vowel, consonant suffix, reflexive imperative, initial vowel"],
            ["language_rules_gr4.py", "Grammar Rules 4.docx",
             "Enumeratives, demonstratives, copula ni-/n-, modal -ta?/-ti, dara, ka particle, fractions, verb-noun derivation, kinship terms"],
            ["language_rules_gr5.py", "Grammar Rules 5.docx (Ch.5,6,7)",
             "Locative prefixes, locative demonstratives, adverbial suffixes, locative possessives, copula+locative, dara+locative, ho+enumerative, objectival concord, noun classes 1a/2a/9a/10a, colour names, negative nouns, professional nouns, augmentatives"],
        ]
    )
    doc.add_paragraph()

    # ── Pipeline order ────────────────────────────────────────────────────────
    heading(doc, "2. Post-Processing Pipeline Order", 1)
    para(doc,
        "Rules are applied in this exact sequence inside _postprocess_lunyoro(). "
        "Order matters — for example, nasal assimilation must run before the R/L rule "
        "because 'nr→nd' would otherwise be incorrectly converted.", size=10)
    table(doc,
        ["Step", "Function", "Module", "What it does"],
        [
            ["1", "apply_nasal_assimilation()", "language_rules.py", "nb→mb, np→mp, nr→nd, nl→nd"],
            ["2", "apply_ni_prefix_change()",   "language_rules.py", "ni+u-class concord → nu (nimugenda→numugenda)"],
            ["3", "apply_consonant_suffix_mutations()", "language_rules.py", "r/t/j/nd/nt + -ire/-i/-ya mutations"],
            ["4", "apply_reflexive_imperative_correction()", "language_rules.py", "okwesereka → weesereke"],
            ["5", "apply_initial_vowel_rule()", "language_rules.py", "Prefix-based initial vowel correction"],
            ["6", "apply_semi_vowel_substitution()", "language_rules.py", "i→y, u→w at prefix boundaries"],
            ["7", "apply_particle_elision()",   "language_rules.py", "na ente→n'ente, habwa okugonza→habw'okugonza"],
            ["8", "apply_rl_rule_to_text()",    "language_rules.py", "L→R except adjacent to e/i"],
            ["9", "apply_gr4_rules()",          "language_rules_gr4.py", "Enumeratives, copula, ka, kinship, modal -ta?"],
            ["10","apply_gr5_rules()",          "language_rules_gr5.py", "Copula+locative, adverbial suffixes"],
            ["11","_normalise_dialect()",       "translate.py", "Rutooro→Runyoro dialect normalisation"],
        ]
    )
    doc.add_paragraph()

    # ── Section 3: Core orthographic rules ───────────────────────────────────
    heading(doc, "3. Core Orthographic Rules (language_rules.py)", 1)

    heading(doc, "3.1  R / L Rule", 2, TEAL)
    para(doc,
        "Source: Runyoro-Rutooro Orthography Guide (1995), Grammar Ch.2\n"
        "R is the dominant consonant. L is only used immediately before or after "
        "the vowels 'e' or 'i'. In all other positions R replaces L.", size=10)
    table(doc, ["Input", "Output", "Reason"],
        [["omulimi",  "omulimi",  "L kept — adjacent to 'i'"],
         ["omulimu",  "omulimo",  "L kept — adjacent to 'i'; final u unchanged"],
         ["okulala",  "okulara",  "Second L→R — not adjacent to e/i"],
         ["okurala",  "okurara",  "L→R — not adjacent to e/i"],
         ["eliiso",   "eliiso",   "L kept — adjacent to 'i'"],
         ["oluganda", "oruganda", "L→R — not adjacent to e/i"]])
    doc.add_paragraph()

    heading(doc, "3.2  Nasal Assimilation", 2, TEAL)
    para(doc,
        "Source: Grammar Ch.2 — Sound change in consonants\n"
        "When n precedes certain consonants, it assimilates to match the place "
        "of articulation of the following consonant.", size=10)
    table(doc, ["Cluster", "Result", "Example"],
        [["nb", "mb", "enbwa → embwa (dog)"],
         ["np", "mp", "enpaka → empaka (cat)"],
         ["nm", "mm", "onmara → ommara"],
         ["nr", "nd", "onruga → ondruga (Meinhof's rule)"],
         ["nl", "nd", "onlimi → ondimi"]])
    doc.add_paragraph()

    heading(doc, "3.3  ni- Prefix Change Before u-Class Concords", 2, TEAL)
    para(doc,
        "Source: Grammar Ch.13 — Present imperfect tense\n"
        "The present imperfect marker ni- changes to nu- when followed by "
        "a u-class concordial prefix (mu-, gu-, ru-, bu-, kw-).", size=10)
    table(doc, ["Input", "Output", "Class"],
        [["nimugenda",  "numugenda",  "Class 1 (mu- concord)"],
         ["niguteera",  "nuguteera",  "Class 3 (gu- concord)"],
         ["niruteera",  "nuruteera",  "Class 11 (ru- concord)"],
         ["nibuteera",  "nubuteera",  "Class 14 (bu- concord)"],
         ["nikwenda",   "nukwenda",   "Class 15 (kw- concord)"]])
    doc.add_paragraph()

    heading(doc, "3.4  Consonant + Suffix Mutations", 2, TEAL)
    para(doc,
        "Source: Grammar Ch.2 — Sound change in verb stems\n"
        "When verb stems ending in r, t, j, nd, or nt take the suffixes "
        "-ire, -i, or -ya, the final consonant mutates.", size=10)
    table(doc, ["Stem ending", "Suffix", "Result", "Example"],
        [["r",  "-ire", "-zire", "okubara → okubazire"],
         ["r",  "-i",   "-zi",   "okubara → okubazi"],
         ["r",  "-ya",  "-za",   "okubara → okubaza"],
         ["t",  "-ire", "-sire", "okukata → okukasire"],
         ["t",  "-i",   "-si",   "okukata → okukasi"],
         ["j",  "-ire", "-zire", "okuhiija → okuhiizire"],
         ["nd", "-ire", "-nzire","okusinda → okusinzire"],
         ["nt", "-ire", "-nsire","okusinta → okusinsire"]])
    doc.add_paragraph()

    heading(doc, "3.5  Semi-Vowel Substitution", 2, TEAL)
    para(doc,
        "Source: Grammar Ch.2 — Vowel/consonant alternation at prefix boundaries\n"
        "At prefix boundaries, the vowels i and u become the semi-vowels y and w "
        "respectively when followed by another vowel.", size=10)
    example_box(doc, "oku + ija → okwija", "to come", "u→w before vowel")
    example_box(doc, "eki + ija → ekyija", "this one coming", "i→y before vowel")
    example_box(doc, "obu + ija → obwija", "the coming", "u→w before vowel")
    doc.add_paragraph()

    heading(doc, "3.6  Apostrophe / Particle Elision", 2, TEAL)
    para(doc,
        "Source: Orthography Guide — Elision rules\n"
        "When certain particles (na, habwa, obu, etc.) precede a vowel-initial word, "
        "the final vowel of the particle is elided and replaced with an apostrophe.", size=10)
    table(doc, ["Before elision", "After elision", "Meaning"],
        [["na ente",          "n'ente",          "with/and a cow"],
         ["na omuntu",        "n'omuntu",        "with/and a person"],
         ["habwa okugonza",   "habw'okugonza",   "because of loving"],
         ["obu aija",         "obw'aija",        "when he comes"],
         ["aha omuntu",       "ah'omuntu",       "at the person"]])
    doc.add_paragraph()

    heading(doc, "3.7  Reflexive Imperative Correction", 2, TEAL)
    para(doc,
        "Source: Grammar Ch.4 — Reflexive verbs\n"
        "Reflexive verb imperatives use wee- (singular) or mwe- (plural) "
        "instead of the full okw- infinitive prefix.", size=10)
    example_box(doc, "okwesereka → weesereke", "hide yourself (sg)", "reflexive imperative")
    example_box(doc, "okwegatta → weegatte", "join yourself (sg)")
    example_box(doc, "okwesereka → mwesereke", "hide yourselves (pl)", "plural reflexive")
    doc.add_paragraph()

    # ── Section 4: Grammar Rules 4 ────────────────────────────────────────────
    heading(doc, "4. Grammar Rules 4 (language_rules_gr4.py)", 1)

    heading(doc, "4.1  Copula ni- / n- Distribution", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Particles ni- and n-\n"
        "The copula (is/are) has two forms. ni- is used before pronouns and "
        "question words. n- is used before demonstratives. Before vowel-initial "
        "words, ni- elides to n' via apostrophe.", size=10)
    table(doc, ["Context", "Form", "Example", "Meaning"],
        [["Before self-standing pronoun", "ni-",  "niinyowe",    "it is I"],
         ["Before 2sg pronoun",           "ni-",  "niiwe",       "it is you"],
         ["Before 3sg pronoun",           "nu-",  "nuwe",        "it is he/she"],
         ["Before 1pl pronoun",           "ni-",  "niitwe",      "it is we"],
         ["Before vowel-initial noun",    "n'",   "n'omuntu",    "it is a person"],
         ["Before vowel-initial noun",    "n'",   "n'ente",      "it is a cow"],
         ["Before near demonstrative",    "n-",   "ngunu (cl.3)","it is this one"],
         ["Before far demonstrative",     "n-",   "nguli (cl.3)","it is that one"],
         ["Merged error corrected",       "n'",   "nomuntu→n'omuntu","common MT error"]])
    doc.add_paragraph()

    heading(doc, "4.2  Kinship Terms with Person Agreement", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Names of Relationship\n"
        "Kinship terms in Runyoro-Rutooro are not standalone nouns — they fuse "
        "with a possessive suffix that agrees with the person. The model often "
        "produces the split form (ise wange) which is corrected to the fused form (isange).", size=10)
    table(doc, ["Relation", "1sg (my)", "2sg (your)", "3sg (his/her)", "1pl (our)"],
        [["father",      "isange",       "isaawe",       "ise",         "isiitwe"],
         ["mother",      "nyinange",     "nyinawe",      "nyina",       "nyinenitu"],
         ["grandfather", "isenkurwange", "isenkurwawe",  "isenkuru",    "isenkurwitwe"],
         ["grandmother", "nyinenkurwange","nyinenkurwawe","nyinenkuru",  "nyinenkurwitwe"]])
    para(doc, "Common MT errors corrected:", bold=True, size=9)
    table(doc, ["MT output (wrong)", "Corrected form", "Relation"],
        [["ise wange",    "isange",    "my father"],
         ["nyina wawe",   "nyinawe",   "your mother"],
         ["isenyowe",     "isange",    "my father (merged error)"],
         ["taata wange",  "isange",    "my father (Swahili loanword)"],
         ["nyinawange",   "nyinange",  "my mother (merged error)"]])
    doc.add_paragraph()

    heading(doc, "4.3  Enumerative Pronouns", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Pronominal Roots\n"
        "Enumerative pronouns express 'alone', 'all', 'themselves', and 'both'. "
        "They are formed by fusing the self-standing pronoun with an enumerative suffix.", size=10)
    table(doc, ["Person", "alone/only (-enka)", "all (-ena)", "themselves (-enyini)", "both (-embi)"],
        [["I",       "nyenka",    "nyeena",   "nyeenyini",  "—"],
         ["you sg",  "wenka",     "weena",    "weenyini",   "—"],
         ["he/she",  "wenka",     "weena",    "weenyini",   "—"],
         ["we",      "twendeka",  "tweena",   "tweenyini",  "twembi"],
         ["you pl",  "mwenka",    "mweena",   "mweenyini",  "mwembi"],
         ["they",    "bonka",     "boona",    "boonyini",   "bombi"]])
    doc.add_paragraph()

    heading(doc, "4.4  Modal Particles: -ota? and -ti", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Modal Particles\n"
        "-ota? is used in greetings to ask 'how?'. -ti introduces reported speech "
        "or describes how something sounds/looks.", size=10)
    table(doc, ["Form", "Usage", "Example", "Meaning"],
        [["-ota?",  "greeting / how?",    "Oraire ota?",         "Good morning / How did you sleep?"],
         ["-ota?",  "state question",     "Oroho ota?",          "How are you?"],
         ["-ti",    "like this",          "Balima bati:",        "They dig like this:"],
         ["-ti",    "reported speech",    "Tukabagambira tuti,", "We told them,"]])
    doc.add_paragraph()

    heading(doc, "4.5  Ka Particle (Emphatic and Permissive)", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Particle ka\n"
        "Ka has two distinct uses: emphatic (the very X) and permissive (let X do Y).", size=10)
    table(doc, ["Use", "Form", "Example", "Meaning"],
        [["Emphatic", "ka + noun",       "ka muntu",    "the very person"],
         ["Emphatic", "ka + pronoun",    "ka niiwe",    "it is really you"],
         ["Emphatic", "ka + near dem",   "ka ngunu",    "here he/she truly is"],
         ["Permissive","ka + subj. verb","ka tugende",  "let us go"],
         ["Permissive","ka + subj. verb","ka ngende",   "let me go"],
         ["Permissive","ka + subj. verb","ka bagende",  "let them go"]])
    doc.add_paragraph()

    heading(doc, "4.6  Dara Presentative Word", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Word dara\n"
        "Dara is a presentative word meaning 'here is / there is'. "
        "It fuses with pronouns and noun-class concords.", size=10)
    table(doc, ["Form", "Meaning"],
        [["daranyowe", "here I am"],
         ["darawe",    "here he/she is"],
         ["darabo",    "here they are"],
         ["daraitwe",  "here we are"],
         ["daroonu",   "here it is (near, cl.1)"],
         ["darakyo",   "here it is (cl.7)"],
         ["darabwo",   "here it is (cl.14)"]])
    doc.add_paragraph()

    heading(doc, "4.7  Verb-to-Noun Derivation", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 4.docx — Noun derivation from verbs\n"
        "Runyoro-Rutooro derives nouns from verb stems using class prefixes and suffixes. "
        "Consonant mutations apply before the -i suffix (r→z, t→s, j→z).", size=10)
    table(doc, ["Type", "Pattern", "Example verb", "Derived noun", "Meaning"],
        [["Agent (cl.1)",  "omu- + root + -i",  "okulima",  "omulimi",  "cultivator"],
         ["Agent (cl.1)",  "omu- + root + -i",  "okubara",  "omubazi",  "carpenter (r→z)"],
         ["Action (cl.3)", "omu- + root + -o",  "okulima",  "omulimo",  "work/digging"],
         ["Action (cl.3)", "omu- + root + -o",  "okuzaana", "omuzaano", "play"],
         ["Method (cl.9)", "en-/em- + stem",    "okulima",  "endima",   "method of digging"],
         ["Method (cl.9)", "en-/em- + stem",    "okubaza",  "embaza",   "method of carpentry"]])
    doc.add_paragraph()

    # ── Section 5: Grammar Rules 5 ────────────────────────────────────────────
    heading(doc, "5. Grammar Rules 5 (language_rules_gr5.py)", 1)

    heading(doc, "5.1  Locative Adverbial Prefixes", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Chapter 5\n"
        "Locative meaning is expressed by prefixing a noun with a locative prefix. "
        "The prefix chosen depends on the spatial relationship.", size=10)
    table(doc, ["Prefix", "Meaning", "Example", "Translation"],
        [["omu-/omw-", "in/inside (enclosed space)", "omunsi",    "in the world"],
         ["omu-/omw-", "in/inside (before vowel)",   "omwiguru",  "in heaven/sky"],
         ["ha-",       "on/at (surface or general)", "hameeza",   "on the table"],
         ["ha-",       "on/at",                      "hansi",     "on the ground"],
         ["ku-",       "to/towards/at",              "kunu",      "this way/side"],
         ["owa-",      "to a person's place",        "owaitu",    "at our home"],
         ["omba",      "to an area belonging to",    "omba so",   "to your father's area"]])
    doc.add_paragraph()

    heading(doc, "5.2  Locative Demonstratives", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — In Combination with Demonstratives\n"
        "Locative prefixes fuse with the near (-nu) or far (-li) demonstrative roots.", size=10)
    table(doc, ["Form", "Prefix + root", "Meaning"],
        [["munu", "omu + nu", "in here (near)"],
         ["muli", "omu + li", "in there (far)"],
         ["hanu", "ha + nu",  "here (near)"],
         ["hali", "ha + li",  "there (far)"],
         ["kunu", "ku + nu",  "this way/side (near)"],
         ["kuli", "ku + li",  "that way/side (far)"]])
    doc.add_paragraph()

    heading(doc, "5.3  Adverbial Suffixes -mu / -ho / -yo", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Locatives as Adverbial Suffixes\n"
        "When a verb is used with a locative noun, the verb takes an adverbial suffix "
        "that matches the locative prefix of the noun. This is a common MT error.", size=10)
    table(doc, ["Locative prefix", "Adverbial suffix", "Example", "Meaning"],
        [["omu-/omw-", "-mu", "Taahamu",          "get in"],
         ["omu-/omw-", "-mu", "harumu amaizi",    "there is water in here"],
         ["ha-",       "-ho", "Taaho",            "take away"],
         ["ha-",       "-ho", "Rugaho",           "get away"],
         ["ha-",       "-ho", "heemeriireho",     "standing on it"],
         ["owa-/ku-",  "-yo", "gendayo owaitu",   "go to our home"],
         ["owa-/ku-",  "-yo", "haraireyo",        "slept at home"]])
    para(doc, "MT errors corrected by apply_adverbial_suffix_correction():", bold=True, size=9)
    table(doc, ["MT output (wrong)", "Corrected", "Rule applied"],
        [["genda owaitu",   "gendayo owaitu",  "-yo suffix added for owa- locative"],
         ["ikara hansi",    "ikaraho hansi",   "-ho suffix added for ha- locative"],
         ["ikara omukibira","ikaramu omukibira","-mu suffix added for omu- locative"]])
    doc.add_paragraph()

    heading(doc, "5.4  Locative Possessives", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — In Combination with Possessives\n"
        "Two locative possessive series exist: omwa- (in someone's house) "
        "and owa- (to/at someone's home).", size=10)
    table(doc, ["Person", "omwa- (in house)", "owa- (to/at home)"],
        [["1sg (my)",    "omwange",  "owange"],
         ["2sg (your)",  "omwawe",   "owaawe"],
         ["3sg (his/her)","omwe",    "owe"],
         ["1pl (our)",   "omwaitu",  "owaitu"],
         ["2pl (your pl)","omwanyu", "owaanyu"],
         ["3pl (their)", "omwabu",   "owaabu"]])
    doc.add_paragraph()

    heading(doc, "5.5  Copula ni- + Locative Adverbials", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — The Copula ni- with adverbials\n"
        "When the copula ni- combines with locative adverbials, the two fuse into "
        "a single word. The model often produces the split form which is corrected.", size=10)
    table(doc, ["Split form (wrong)", "Fused form (correct)", "Meaning"],
        [["ni hanu",  "nihanu",  "it is here"],
         ["ni hali",  "nihali",  "it is there"],
         ["ni kunu",  "nukunu",  "it is this way"],
         ["ni kuli",  "nukuli",  "it is that way"],
         ["ni munu",  "nimunu",  "it is in here"],
         ["ni muli",  "nimuli",  "it is in there"],
         ["ni mwo",   "numwo",   "it is in it"],
         ["ni ho",    "nuho",    "it is on it"],
         ["ni oku",   "nooku",   "it is that direction"],
         ["ni aho",   "naaho",   "it is on there"]])
    doc.add_paragraph()

    heading(doc, "5.6  ho + Enumerative Roots", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — The Adverbials mwo and ho\n"
        "The locative adverbial ho (on it) combines with enumerative roots "
        "to form compound adverbials.", size=10)
    table(doc, ["Form", "Meaning", "Example"],
        [["hoona",    "everywhere / all over",  "Ruhanga ali omwiguru, n'omunsi na buli hantu hoona."],
         ["honka",    "only there",             "Nkarora omusanduuko honka."],
         ["hombi",    "both sides/places",      "Hombi, omunda n'aheeru y'orugoye runu nihasisana."],
         ["hoonyini", "the very spot",          "Ho hoonyini ntahikeho."]])
    doc.add_paragraph()

    heading(doc, "5.7  Objectival Concord (Reversed-Object Sentences)", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Chapter 6, Object\n"
        "When the object is fronted (moved to the beginning of the sentence), "
        "the verb takes an objectival concord infix before the stem.", size=10)
    table(doc, ["Noun class", "Objectival concord", "Example"],
        [["1 (person)",  "mu", "omusiri omukazi agulimire — the woman dug the garden"],
         ["3 (tree)",    "gu", "omuti omukazi agutemire — the woman cut the tree"],
         ["7 (thing)",   "ki", "ekitabo omwana akitwaire — the child took the book"],
         ["9 (animal)",  "yi", "ente omuhuma ayinyire — the herdsman drank the milk"],
         ["10 (plural)", "zi", "ente abahuma baziinyire — the herdsmen drank the cows' milk"]])
    doc.add_paragraph()

    heading(doc, "5.8  Special Noun Classes (1a, 2a, 9a, 10a)", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Chapter 7\n"
        "Four special noun classes handle proper names, foreign words, and colours.", size=10)
    table(doc, ["Class", "Description", "Prefix", "Examples"],
        [["1a",  "Proper names, empaako, animal names (no prefix)", "—",    "Abbooki, mukaaka, warujojo"],
         ["2a",  "Plural of class 1a",                             "baa-", "Baabbooki, Baamukaaka"],
         ["9a",  "Foreign words, colours, place names (no prefix)", "—",   "motoka, kinyansi, Buganda"],
         ["10a", "Plural of class 9a",                             "zaa-", "zaamotoka, zaakinyansi"]])
    doc.add_paragraph()

    heading(doc, "5.9  Colour Names (Class 9a)", 2, TEAL)
    table(doc, ["English", "Runyoro-Rutooro", "Literal reference"],
        [["green",       "kinyansi",  "like grass"],
         ["white",       "kyeru",     "like a white cow/hen"],
         ["black",       "kikara",    "like a black cow/hen"],
         ["brown",       "kitaka",    "like soil"],
         ["red/reddish", "kigaaja",   "like a reddish cow"],
         ["grey",        "kibuubi",   "like a greyish cow"],
         ["dark brown",  "kisiina",   "like a dark brown cow"],
         ["yellow",      "kyenju",    "like a ripe banana"],
         ["blue",        "bbururu",   "—"],
         ["purple",      "kihuukya",  "like ehuukya berries"]])
    doc.add_paragraph()

    heading(doc, "5.10  Negative Nouns (omu-ta- prefix)", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Class 1 Nouns\n"
        "A person who habitually does NOT do something is described by "
        "prefixing omu-ta- to the verb stem.", size=10)
    table(doc, ["Form", "Verb stem", "Meaning"],
        [["omutaseka",     "seka (to laugh)",    "gloomy person / one who does not laugh"],
         ["omutooga",      "tooga (to bathe)",   "dirty person / one who does not bathe"],
         ["omutagambwaho", "gambwaho",           "touchy person / easily offended"]])
    doc.add_paragraph()

    heading(doc, "5.11  Class 9 Professional Nouns (en-/em- prefix)", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Class 9 Nouns\n"
        "A professional or habitual doer is derived by prefixing en- (or em- before b/p) "
        "to the verb stem. These differ from the class 1 agent noun (omulimi).", size=10)
    table(doc, ["Form", "Meaning", "Compare to"],
        [["endimi",    "professional cultivator",  "omulimi (cultivator)"],
         ["ensuubuzi", "professional trader",      "omusuubuzi (trader)"],
         ["encwangya", "incurable liar",           "omucwangya (liar)"],
         ["enfaakati", "permanent widow/widower",  "mufaakati"]])
    doc.add_paragraph()

    heading(doc, "5.12  Augmentative / Pejorative Prefix Substitution", 2, TEAL)
    para(doc,
        "Source: Grammar Rules 5.docx — Class 5 and Class 7 Emotional Values\n"
        "Emotional nuance (contempt, affection, magnitude) is expressed by substituting "
        "the normal class prefix with class 5 (eri-/i-) or class 7 (eki-) prefix.", size=10)
    table(doc, ["Form", "Base noun", "Emotional value", "Meaning"],
        [["isaija",   "omusaija (man)",  "Class 5 — pejorative",  "big/disrespectful man"],
         ["isigazi",  "omusigazi (youth)","Class 5 — pejorative", "youth acting badly"],
         ["eryana",   "omwana (child)",  "Class 5 — pejorative",  "insolent child"],
         ["ekisaija", "omusaija (man)",  "Class 7 — contempt",    "clumsy/contemptible man"],
         ["ekiiru",   "omwiru (serf)",   "Class 7 — affection",   "dear poor man"]])
    doc.add_paragraph()

    # ── Section 6: Noun class system ─────────────────────────────────────────
    heading(doc, "6. Noun Class System", 1)
    para(doc,
        "Source: Grammar Ch.7 — The noun class system\n"
        "Runyoro-Rutooro has 15 noun classes (plus 4 special classes). "
        "Every noun belongs to a class, and all agreeing words (verbs, adjectives, "
        "pronouns, demonstratives) must carry the matching concordial prefix.", size=10)
    table(doc,
        ["Class", "Singular prefix", "Plural class", "Plural prefix", "Typical content"],
        [["1",   "omu-/omw-", "2",  "aba-/ab-",  "persons (singular)"],
         ["2",   "aba-/ab-",  "1",  "omu-/omw-", "persons (plural)"],
         ["1a",  "— (none)",  "2a", "baa-",       "proper names, empaako, animal names"],
         ["3",   "omu-/omw-", "4",  "emi-",       "trees, plants, body parts (sg)"],
         ["4",   "emi-",      "3",  "omu-",       "plural of class 3"],
         ["5",   "eri-/ery-", "6",  "ama-",       "augmentatives, some body parts (sg)"],
         ["6",   "ama-",      "5",  "eri-",       "plural of cl.5; also cl.9/11/14/15"],
         ["7",   "eki-/eky-", "8",  "ebi-/eby-",  "things, abstracts, diminutives (sg)"],
         ["8",   "ebi-/eby-", "7",  "eki-/eky-",  "plural of class 7"],
         ["9",   "en-/em-",   "10", "en-/em-",    "animals, foreign words (sg)"],
         ["10",  "en-/em-",   "9",  "en-/em-",    "plural of class 9"],
         ["9a",  "— (none)",  "10a","zaa-",        "foreign words, colours, place names"],
         ["11",  "oru-/orw-", "10", "en-",         "long/thin objects, languages, abstract"],
         ["12",  "aka-/akw-", "13", "utu-",        "diminutives (singular)"],
         ["13",  "utu-/utw-", "12", "aka-",        "diminutives (plural)"],
         ["14",  "obu-/obw-", "6",  "ama-",        "abstract nouns, mass nouns"],
         ["15",  "oku-/okw-", "6",  "ama-",        "verbal infinitives, body parts"]])
    doc.add_paragraph()

    # ── Section 7: Tense system ───────────────────────────────────────────────
    heading(doc, "7. Tense and Aspect System", 1)
    para(doc,
        "Source: Grammar Ch.13, Ch.15\n"
        "Tense and aspect are marked morphologically on the verb. "
        "The tense marker is inserted between the subject prefix and the verb stem.", size=10)
    table(doc, ["Tense/Aspect", "Marker", "Example", "Meaning"],
        [["Present imperfect",  "ni-",    "nigenda",      "is going"],
         ["Present perfect",   "-ire",   "agenzire",     "has gone"],
         ["Recent past",       "a-",     "nayara",       "I just made the bed"],
         ["Remote past",       "ka-",    "nkaara",       "I made the bed (remote)"],
         ["Future immediate",  "ra-",    "ndaayara",     "I shall make the bed"],
         ["Future remote",     "raa-",   "turaayara",    "we shall make the bed"],
         ["Imperative sg",     "stem-a", "genda",        "go! (singular)"],
         ["Imperative pl",     "mu-stem-e","mugende",    "go! (plural)"],
         ["Negative present",  "ti-ni-", "tinigenda",    "is not going"],
         ["Negative perfect",  "tinka-", "tinkagenzire", "has not gone"],
         ["Conditional",       "-ku-",   "obaire okukora","if/when doing"]])
    doc.add_paragraph()

    # ── Section 8: Dialect normalisation ─────────────────────────────────────
    heading(doc, "8. Dialect Normalisation (Rutooro → Runyoro)", 1)
    para(doc,
        "Source: translate.py — _normalise_dialect()\n"
        "The translator targets standard Runyoro forms. When the model produces "
        "Rutooro dialect variants, they are normalised to the Runyoro standard.", size=10)
    table(doc, ["Rutooro form", "Runyoro standard", "Meaning"],
        [["kiro kinu",      "leero",           "today"],
         ["kiro ekindi",    "leero",           "today"],
         ["kyakabizi",      "n'Orwokasatu",    "Tuesday"],
         ["kya kabizi",     "n'Orwokasatu",    "Tuesday"],
         ["kya kasatu",     "n'Orwokasatu",    "Tuesday"],
         ["kya kana",       "n'Orwokana",      "Thursday"],
         ["kya kataano",    "n'Orwokataano",   "Friday"],
         ["kya mukaaga",    "n'Orwomukaaga",   "Saturday"],
         ["kya sande",      "n'Orwosande",     "Sunday"],
         ["kya banza",      "n'Orwobanza",     "Monday"],
         ["kiro",           "leero",           "today (standalone)"]])
    doc.add_paragraph()

    # ── Section 9: Where rules are applied ───────────────────────────────────
    heading(doc, "9. Where Rules Are Applied in the System", 1)
    para(doc,
        "Rules are applied at different points in the translation pipeline "
        "depending on whether they affect input normalisation or output correction.", size=10)
    table(doc, ["Stage", "Function", "Rules applied"],
        [["Input (lun→en)",  "_preprocess_lunyoro_input()",
          "Nasal assimilation only — normalises input before feeding to model"],
         ["Output (en→lun)", "_postprocess_lunyoro()",
          "Full pipeline: all 11 steps in order (see Section 2)"],
         ["Chat system prompt","_GRAMMAR_CONTEXT_CACHE",
          "Grammar context injected into Ollama/Qwen LLM for chat responses"],
         ["Word lookup",     "lookup_word()",
          "R/L rule applied to MT output for dictionary lookups"],
         ["Document summary","summarize_pdf()",
          "Gr4 rules applied to Lunyoro output of sentence-by-sentence translation"]])
    doc.add_paragraph()

    # ── Section 10: Summary of all corrections ────────────────────────────────
    heading(doc, "10. Summary: All Active Corrections", 1)
    para(doc,
        "This table lists every regex-based correction that runs automatically "
        "on every en→lun translation output.", size=10)
    table(doc, ["Rule", "Pattern corrected", "Function"],
        [["Nasal assimilation",    "nb→mb, np→mp, nr→nd, nl→nd",          "apply_nasal_assimilation()"],
         ["ni-prefix change",      "nimu→numu, nigu→nugu, niru→nuru",     "apply_ni_prefix_change()"],
         ["Consonant mutations",   "r/t/j/nd/nt + -ire/-i/-ya",           "apply_consonant_suffix_mutations()"],
         ["Reflexive imperative",  "okweX → weeX (sg), mweX (pl)",        "apply_reflexive_imperative_correction()"],
         ["Semi-vowel",            "i→y, u→w at prefix boundaries",       "apply_semi_vowel_substitution()"],
         ["Particle elision",      "na+vowel→n', habwa+vowel→habw'",      "apply_particle_elision()"],
         ["R/L rule",              "L→R except adjacent to e/i",          "apply_rl_rule_to_text()"],
         ["Enumerative pronouns",  "yenka→wenka, yoona→boona",            "apply_enumerative_correction()"],
         ["Copula ni-/n-",         "ni omuntu→n'omuntu, ni nyowe→niinyowe","apply_copula_to_text()"],
         ["Ka emphatic",           "ka + demonstrative preservation",      "apply_ka_emphatic()"],
         ["Kinship terms",         "ise wange→isange, nyina wawe→nyinawe","apply_kinship_correction()"],
         ["Modal -ta? greetings",  "Oraire ota? preserved",               "apply_modal_ta_greeting()"],
         ["Copula + locative",     "ni hanu→nihanu, ni mwo→numwo",        "apply_copula_locative_correction()"],
         ["Adverbial suffixes",    "genda owaitu→gendayo owaitu",         "apply_adverbial_suffix_correction()"],
         ["Dialect normalisation", "kiro→leero, kya kataano→n'Orwokataano","_normalise_dialect()"]])

    doc.add_paragraph()
    para(doc,
        f"Document generated: {datetime.now().strftime('%d %B %Y, %H:%M')}  |  "
        "Sources: A Grammar of Runyoro-Rutooro (Ch.2,4,7,13,15,16), "
        "Runyoro-Rutooro Orthography Guide (1995), "
        "Grammar Rules 4.docx, Grammar Rules 5.docx",
        italic=True, size=8, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
